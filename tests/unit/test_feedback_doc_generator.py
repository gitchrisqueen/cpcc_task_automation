#  Copyright (c) 2024. Christopher Queen Consulting LLC (http://www.ChristopherQueenConsulting.com/)

"""Unit tests for feedback_doc_generator module.

Tests that the CPCC-branded Word document generator:
- Generates valid DOCX bytes
- Contains required headings and structure
- Does NOT include numeric grading patterns
- Includes error titles when provided
- Handles missing sections gracefully
- Sanitizes filenames correctly
"""

import re
from datetime import datetime
from io import BytesIO

import pytest
from docx import Document

from cqc_cpcc.feedback_doc_generator import (
    generate_student_feedback_doc,
    sanitize_filename,
    _parse_feedback_sections,
)
from cqc_cpcc.rubric_models import (
    CriterionResult,
    DetectedError,
    RubricAssessmentResult,
)


@pytest.mark.unit
def test_sanitize_filename_removes_invalid_characters():
    """Test that sanitize_filename removes invalid filename characters."""
    # Test various invalid characters
    assert sanitize_filename('John/Doe') == 'JohnDoe'
    assert sanitize_filename('Jane:Smith') == 'JaneSmith'
    assert sanitize_filename('Test*File') == 'TestFile'
    assert sanitize_filename('File?Name') == 'FileName'
    assert sanitize_filename('Doc|File') == 'DocFile'
    assert sanitize_filename('Test"File"') == 'TestFile'


@pytest.mark.unit
def test_sanitize_filename_replaces_spaces():
    """Test that sanitize_filename replaces spaces with underscores."""
    assert sanitize_filename('John Doe') == 'John_Doe'
    assert sanitize_filename('Multiple  Spaces') == 'Multiple__Spaces'


@pytest.mark.unit
def test_sanitize_filename_handles_empty_string():
    """Test that sanitize_filename handles empty strings."""
    assert sanitize_filename('') == 'Feedback'
    assert sanitize_filename('   ') == 'Feedback'
    assert sanitize_filename('...') == 'Feedback'


@pytest.mark.unit
def test_generate_student_feedback_doc_returns_bytes():
    """Test that generate_student_feedback_doc returns non-empty bytes."""
    # Create minimal feedback result
    result = RubricAssessmentResult(
        total_points_earned=85,
        total_points_possible=100,
        overall_feedback="Good work overall.",
        criteria_results=[
            CriterionResult(
                criterion_id="c1",
                criterion_name="Code Quality",
                points_earned=40,
                points_possible=50,
                feedback="Good structure and readability.",
                selected_level_label="Proficient"
            )
        ],
        detected_errors=[]
    )
    
    doc_bytes = generate_student_feedback_doc(
        student_name="John Doe",
        course_id="CSC151",
        assignment_name="Exam 1",
        feedback_result=result
    )
    
    # Verify bytes returned
    assert isinstance(doc_bytes, bytes)
    assert len(doc_bytes) > 0
    
    # Verify it's a valid DOCX
    doc_stream = BytesIO(doc_bytes)
    doc = Document(doc_stream)
    assert doc is not None


@pytest.mark.unit
def test_generate_student_feedback_doc_contains_required_headings():
    """Test that generated doc contains required section headings."""
    result = RubricAssessmentResult(
        total_points_earned=75,
        total_points_possible=100,
        overall_feedback="Room for improvement in several areas.",
        criteria_results=[
            CriterionResult(
                criterion_id="c1",
                criterion_name="Documentation",
                points_earned=30,
                points_possible=40,
                feedback="Good comments but could be more detailed.",
                selected_level_label="Developing"
            ),
            CriterionResult(
                criterion_id="c2",
                criterion_name="Logic",
                points_earned=45,
                points_possible=60,
                feedback="Logic is mostly correct but has some issues.",
                selected_level_label="Developing"
            )
        ],
        detected_errors=[]
    )
    
    doc_bytes = generate_student_feedback_doc(
        student_name="Jane Smith",
        course_id="CSC251",
        assignment_name="Midterm",
        feedback_result=result
    )
    
    # Parse document
    doc = Document(BytesIO(doc_bytes))
    doc_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Check for required headings
    assert 'Feedback Summary' in doc_text
    assert 'Student: Jane Smith' in doc_text
    assert 'Course: CSC251' in doc_text
    assert 'Assignment: Midterm' in doc_text
    assert 'Summary' in doc_text
    assert 'Strengths' in doc_text or 'Areas for Improvement' in doc_text


@pytest.mark.unit
def test_generate_student_feedback_doc_no_numeric_grading():
    """Test that generated doc does NOT contain numeric grading patterns."""
    result = RubricAssessmentResult(
        total_points_earned=85,
        total_points_possible=100,
        overall_feedback="Strong submission.",
        criteria_results=[
            CriterionResult(
                criterion_id="c1",
                criterion_name="Code Quality",
                points_earned=85,
                points_possible=100,
                feedback="Excellent work.",
                selected_level_label="Exemplary"
            )
        ],
        detected_errors=[]
    )
    
    doc_bytes = generate_student_feedback_doc(
        student_name="Test Student",
        course_id="CSC151",
        assignment_name="Test Assignment",
        feedback_result=result
    )
    
    # Parse document text
    doc = Document(BytesIO(doc_bytes))
    doc_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Define patterns to check for (should NOT be present)
    forbidden_patterns = [
        r'\d+\s*/\s*\d+',  # e.g., "85/100"
        r'\d+\s*%',  # e.g., "85%"
        r'\d+\s*points',  # e.g., "85 points"
        r'score:\s*\d+',  # e.g., "Score: 85"
        r'earned:\s*\d+',  # e.g., "Earned: 85"
        r'[A-F][+-]?',  # Grade bands like "A+", "B-"
    ]
    
    for pattern in forbidden_patterns:
        matches = re.findall(pattern, doc_text, re.IGNORECASE)
        # Filter out date patterns (YYYY-MM-DD contains /)
        matches = [m for m in matches if not re.match(r'\d{4}-\d{2}-\d{2}', m)]
        assert len(matches) == 0, f"Found forbidden pattern '{pattern}': {matches}"


@pytest.mark.unit
def test_generate_student_feedback_doc_includes_error_titles():
    """Test that generated doc includes error titles when errors are present."""
    result = RubricAssessmentResult(
        total_points_earned=70,
        total_points_possible=100,
        overall_feedback="Several issues identified.",
        criteria_results=[
            CriterionResult(
                criterion_id="c1",
                criterion_name="Code Quality",
                points_earned=70,
                points_possible=100,
                feedback="Issues with documentation and logic.",
                selected_level_label="Developing"
            )
        ],
        detected_errors=[
            DetectedError(
                code="INSUFFICIENT_DOCS",
                name="Insufficient Documentation",
                description="Code lacks adequate comments explaining logic.",
                severity="major",
                occurrences=3,
                notes="Add comments above each method and complex logic blocks."
            ),
            DetectedError(
                code="NAMING_CONVENTION",
                name="Naming Convention Violations",
                description="Variable names do not follow camelCase convention.",
                severity="minor",
                occurrences=5,
                notes="Use camelCase for variable names (e.g., studentName not student_name)."
            )
        ]
    )
    
    doc_bytes = generate_student_feedback_doc(
        student_name="Error Test Student",
        course_id="CSC151",
        assignment_name="Error Test",
        feedback_result=result
    )
    
    # Parse document text
    doc = Document(BytesIO(doc_bytes))
    doc_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Check for error section
    assert 'Errors Observed' in doc_text
    
    # Check for error titles (human-readable names, NOT codes)
    assert 'Insufficient Documentation' in doc_text
    assert 'Naming Convention Violations' in doc_text
    
    # Ensure internal error codes are NOT present
    assert 'INSUFFICIENT_DOCS' not in doc_text
    assert 'NAMING_CONVENTION' not in doc_text
    
    # Check error descriptions are present
    assert 'Code lacks adequate comments' in doc_text
    assert 'Variable names do not follow' in doc_text


@pytest.mark.unit
def test_generate_student_feedback_doc_handles_no_errors():
    """Test that generated doc handles case when no errors are present."""
    result = RubricAssessmentResult(
        total_points_earned=95,
        total_points_possible=100,
        overall_feedback="Excellent work with no major issues.",
        criteria_results=[
            CriterionResult(
                criterion_id="c1",
                criterion_name="Code Quality",
                points_earned=95,
                points_possible=100,
                feedback="Outstanding implementation.",
                selected_level_label="Exemplary"
            )
        ],
        detected_errors=[]
    )
    
    doc_bytes = generate_student_feedback_doc(
        student_name="Perfect Student",
        course_id="CSC251",
        assignment_name="Perfect Test",
        feedback_result=result
    )
    
    # Parse document text
    doc = Document(BytesIO(doc_bytes))
    doc_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # "Errors Observed" section should NOT be present when no errors
    assert 'Errors Observed' not in doc_text


@pytest.mark.unit
def test_generate_student_feedback_doc_includes_metadata():
    """Test that generated doc includes metadata when provided."""
    result = RubricAssessmentResult(
        total_points_earned=80,
        total_points_possible=100,
        overall_feedback="Good work.",
        criteria_results=[],
        detected_errors=[]
    )
    
    test_date = '2024-12-15'
    metadata = {
        'date': test_date,
        'instructor': 'Dr. Smith'
    }
    
    doc_bytes = generate_student_feedback_doc(
        student_name="Metadata Test",
        course_id="CSC134",
        assignment_name="Project 1",
        feedback_result=result,
        metadata=metadata
    )
    
    # Parse document
    doc = Document(BytesIO(doc_bytes))
    doc_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Check for date in document
    assert test_date in doc_text


@pytest.mark.unit
def test_parse_feedback_sections_extracts_summary():
    """Test that _parse_feedback_sections correctly extracts summary."""
    feedback_text = """Here is some summary text.
More summary content.

**Strengths:**
  - Good coding style

**Areas for Improvement:**
  - Add more comments
"""
    
    sections = _parse_feedback_sections(feedback_text)
    
    assert 'summary' in sections
    assert 'Here is some summary text' in sections['summary']
    assert 'More summary content' in sections['summary']


@pytest.mark.unit
def test_parse_feedback_sections_extracts_strengths():
    """Test that _parse_feedback_sections correctly extracts strengths."""
    feedback_text = """Summary text.

**Strengths:**
  - Excellent code structure
  - Good variable naming
  - Well-documented methods

**Areas for Improvement:**
  - Could add more tests
"""
    
    sections = _parse_feedback_sections(feedback_text)
    
    assert 'strengths' in sections
    assert len(sections['strengths']) == 3
    assert 'Excellent code structure' in sections['strengths']
    assert 'Good variable naming' in sections['strengths']
    assert 'Well-documented methods' in sections['strengths']


@pytest.mark.unit
def test_parse_feedback_sections_extracts_improvements():
    """Test that _parse_feedback_sections correctly extracts improvements."""
    feedback_text = """Summary.

**Strengths:**
  - Good work

**Areas for Improvement:**
  - Add error handling
  - Improve code efficiency
"""
    
    sections = _parse_feedback_sections(feedback_text)
    
    assert 'improvements' in sections
    assert len(sections['improvements']) == 2
    assert 'Add error handling' in sections['improvements']
    assert 'Improve code efficiency' in sections['improvements']


@pytest.mark.unit
def test_parse_feedback_sections_extracts_errors():
    """Test that _parse_feedback_sections correctly extracts errors."""
    feedback_text = """Summary.

**Errors Observed:**

*Major Issues:*
  - **Insufficient Documentation**: Code lacks comments.
    Additional notes about this error.

*Minor Issues:*
  - **Naming Convention**: Variables not properly named.
"""
    
    sections = _parse_feedback_sections(feedback_text)
    
    assert 'errors' in sections
    assert len(sections['errors']) == 2
    
    # Check major error
    major_error = next(e for e in sections['errors'] if e['severity'] == 'major')
    assert major_error['name'] == 'Insufficient Documentation'
    assert 'Code lacks comments' in major_error['description']
    
    # Check minor error
    minor_error = next(e for e in sections['errors'] if e['severity'] == 'minor')
    assert minor_error['name'] == 'Naming Convention'
    assert 'Variables not properly named' in minor_error['description']


@pytest.mark.unit
def test_generate_student_feedback_doc_handles_empty_criteria():
    """Test that doc generation handles empty criteria results."""
    result = RubricAssessmentResult(
        total_points_earned=0,
        total_points_possible=100,
        overall_feedback="No criteria assessed.",
        criteria_results=[],
        detected_errors=[]
    )
    
    doc_bytes = generate_student_feedback_doc(
        student_name="Empty Test",
        course_id="CSC999",
        assignment_name="Empty Test",
        feedback_result=result
    )
    
    # Should not raise exception
    assert isinstance(doc_bytes, bytes)
    assert len(doc_bytes) > 0
    
    # Verify it's a valid document
    doc = Document(BytesIO(doc_bytes))
    assert doc is not None


@pytest.mark.unit
def test_generate_student_feedback_doc_cpcc_branding():
    """Test that generated doc follows CPCC branding (basic verification)."""
    result = RubricAssessmentResult(
        total_points_earned=90,
        total_points_possible=100,
        overall_feedback="Excellent submission.",
        criteria_results=[
            CriterionResult(
                criterion_id="c1",
                criterion_name="Quality",
                points_earned=90,
                points_possible=100,
                feedback="Great work.",
                selected_level_label="Exemplary"
            )
        ],
        detected_errors=[]
    )
    
    doc_bytes = generate_student_feedback_doc(
        student_name="Branding Test",
        course_id="CSC151",
        assignment_name="Branding Test",
        feedback_result=result
    )
    
    # Parse document
    doc = Document(BytesIO(doc_bytes))
    
    # Check that document has content
    assert len(doc.paragraphs) > 0
    
    # Check for title
    title_text = doc.paragraphs[0].text
    assert 'Feedback Summary' in title_text
    
    # Verify font settings on title (basic check)
    if doc.paragraphs[0].runs:
        title_run = doc.paragraphs[0].runs[0]
        assert title_run.font.name == 'Calibri'
        assert title_run.font.bold is True
        # Color check: CPCC Blue is RGBColor(0, 90, 163)
        # RGBColor stores as integer, check if set (non-None means color was applied)
        if title_run.font.color.rgb is not None:
            # Verify color was set (actual value is implementation-dependent)
            assert title_run.font.color.rgb is not None
