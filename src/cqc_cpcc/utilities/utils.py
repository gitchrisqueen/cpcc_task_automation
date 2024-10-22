import os
import os.path
import tempfile
import time
import zipfile
from enum import Enum, StrEnum
from random import randint
from typing import Optional, Annotated, List, Union

import docx
import mammoth
# import markdownify
import pandas as pd
import streamlit
import textract
from bs4 import BeautifulSoup
from docx import Document
from markdownify import markdownify as md
from ordered_set import OrderedSet
from pydantic.v1 import BaseModel, Field, StrictStr, PositiveInt
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC

from cqc_cpcc.utilities.date import get_datetime
from cqc_cpcc.utilities.selenium_util import get_driver_wait, click_element_wait_retry


# from simplify_docx import simplify

# Global Constants
LINE_DASH_COUNT = 33



class Satisfactory(Enum):
    YES = 1
    NO = 0


def are_you_satisfied():
    """Prompts the user to select if they are satisfied or not."""

    enum = Satisfactory

    print("Are you satisfied?")
    for i, member in enumerate(enum):
        print(f"{member.value}: {member.name}")

    default = Satisfactory.YES
    default_value = default.value
    user_input = int(input('Enter your selection [' + str(default_value) + ']: ').strip() or default_value)

    try:
        sf = Satisfactory(user_input)
        print(f"You selected {sf.name}")
        return sf.value == default_value
    except ValueError:
        print("Invalid selection.")
        return are_you_satisfied() == default_value


def first_two_uppercase(string):
    """Returns the first two letters of a string, in uppercase."""
    return string[:2].upper()


def html_table_to_dict(table_text) -> [dict, dict]:
    # logger.info('Table: %s' %table_text)
    table_text = table_text.replace("\n", " ")  # Remove new lines
    # TODO: ??? Add Pipe seperator to <spans> or <p>
    table_text = table_text.replace("</p>", "|</p>")  # This is to be able to split assignment names later
    soup = BeautifulSoup(table_text, 'html.parser')

    rows = []
    headers = []

    for tr in soup.find('table').find('tbody').find_all('tr'):
        row = []

        for th in tr.find_all('th'):
            try:
                headers.append(th.text.strip())
            except:
                continue

        for td in tr.find_all('td'):
            try:
                row.append(td.text.strip())
            except:
                continue
        # logger.info('Row Length : %s' % str(len(row)))
        if 1 < len(row) < 3:
            # Modify the 2nd and 3rd column so its has the same date text
            row_date = get_datetime(row[1].split("-")[0].strip())
            row[1] = row_date.strftime("%Y-%m-%d")
            row.append(row[1])
        if len(row) == 3:
            rows.append(row)

    # logger.info('Table Headers: %s' % headers)
    # logger.info('Table (extracted): %s' % rows)

    df = pd.DataFrame(rows, columns=headers)

    return headers, df.to_dict()


def get_unique_names(list_names: list):
    # insert the list to the set
    list_set = set(list_names)
    # convert the set to the list
    unique_list = (list(list_set))
    # Sort alphabetically
    unique_list.sort()

    return unique_list


def get_unique_names_flip_first_last(list_names: list) -> list:
    unique_names = get_unique_names(list_names)
    names_flipped = map(lambda kv: flip_name(kv), unique_names)
    return list(names_flipped)


def flip_name(full_name: str):
    separator = ','
    name_parts = full_name.split(separator)
    name_parts.reverse()
    return separator.join(name_parts)


"""
def get_html_as_markdown(html: str, code_language='java') -> str:
    # convert html to markdown

    markdown = markdownify.markdownify(html, heading_style="ATX", code_language=code_language)
    return markdown
"""


class ExtendedEnum(StrEnum):

    @classmethod
    def list(cls) -> list:
        return [c.value for c in cls]


class CodeError(BaseModel):
    """Object representing a Code Error"""
    error_type: Optional[
        Annotated[
            Enum,
            Field(description="The type of coding error")
        ]
    ] = None
    code_error_lines: Optional[
        Annotated[
            List[StrictStr],
            Field(description="An array list of the lines of code that are relevant to the coding error")
        ]
    ] = None

    line_numbers_of_error_holder: Optional[
        Annotated[
            List[PositiveInt],
            Field(description="The list of line numbers relevant to this coding error")
        ]
    ] = None

    error_details: Annotated[
        StrictStr,
        Field(description="The details about this coding error")
    ]

    def set_line_numbers_of_error(self, line_numbers_of_errors: List[PositiveInt]):
        self.line_numbers_of_error_holder = line_numbers_of_errors

    @property
    def line_numbers_of_errors(self) -> List:
        if self.line_numbers_of_error_holder is None:
            return []
        else:
            return sorted(set(self.line_numbers_of_error_holder))

    def __str__(self):
        import cqc_cpcc.utilities.env_constants as EC
        lines_string_complete = ""
        if EC.SHOW_ERROR_LINE_NUMBERS and self.line_numbers_of_errors is not None:
            lines_string = ", ".join(map(str, self.line_numbers_of_errors))
            lines_string_complete = f"\n\tOn Line(s) #: {lines_string}"
        error_details_string = "\t" + self.error_details.replace("\n", "\n\t")
        return f"{self.error_type.value}:{lines_string_complete}\n{error_details_string}"


class ErrorHolder(BaseModel):

    def get_combined_errors_by_type(self, code_errors: List[CodeError]) -> List[CodeError]:
        errors = {}
        final_errors = []
        for code_error in code_errors:
            if code_error.error_type not in errors:
                errors[code_error.error_type] = []
            errors[code_error.error_type].append(code_error)

        # print("Errors by Type")
        # pprint(errors)

        for error_type in errors:
            code_errors_by_type = errors[error_type]
            nested_list = [x.line_numbers_of_errors for x in code_errors_by_type]
            # print("Nested List")
            # pprint(nested_list)
            # Use map and lambda to flatten the nested list
            # Use list comprehension to flatten the nested list
            try:
                flattened_list = [number for sublist in nested_list for number in sublist]
            except TypeError as e:
                flattened_list = []

            # print("Flattened List")
            # pprint(flattened_list)
            unique_sorted_list = sorted(set(flattened_list))
            # print("Unique Sorted List")
            # pprint(unique_sorted_list)
            # error_details_list = list(map(lambda x: x.error_details, code_errors_by_type))

            # Combine the error details
            error_details_list = [x.error_details for x in code_errors_by_type]
            unique_error_details_list = list(OrderedSet(error_details_list))
            error_details = "\n".join(unique_error_details_list)

            # Combine the code_error_lines
            try:
                code_error_lines_list = [x.code_error_lines for x in code_errors_by_type]
                code_error_lines_flattened_list = [line for code_error_lines_sublist in code_error_lines_list for line
                                                   in
                                                   code_error_lines_sublist]
            except TypeError as e:
                code_error_lines_flattened_list = []

            final_errors.append(
                CodeError(error_type=error_type,
                          line_numbers_of_error=unique_sorted_list,
                          # TODO: Do we need this since the line numbers come after from code???
                          error_details=error_details,
                          code_error_lines=code_error_lines_flattened_list))

        # print("Final Errors by Type Combined")
        # pprint(final_errors)

        return final_errors


def wrap_code_in_markdown_backticks(code: str, code_type: str = "java") -> str:
    backticks = "`"
    code_fence = backticks * 3
    # See if code fence exist in code already
    while code_fence in code:
        code_fence += backticks

    # Add double-backslash to code
    # code = code.replace('\\','\\\\')

    prefix = code_fence + code_type + "\n"
    suffix = "\n" + code_fence
    return prefix + code.strip() + suffix


def merge_lists(list1, list2):
    """Merges two lists, where one might be None.

    Args:
      list1: The first list.
      list2: The second list.

    Returns:
      A merged list.
    """
    if list1 is None and list2 is None:
        return None
    elif list1 is None:
        return list2
    elif list2 is None:
        return list1
    else:
        return list1 + list2


def convert_tables_to_json_in_tmp__file(doc: Document) -> str:
    for table in doc.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data)

        # Remove the table
        t = table._element
        parent = t.getparent()
        parent.remove(t)

        # Add new json string to the parent in its place
        doc.add_paragraph(df.to_json(orient="records"))

        # Clear the table reference
        t._t = t._element = None

    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)

    return temp_file.name


@streamlit.cache_data
def convert_content_to_markdown(content: str) -> str:
    return md(content)


@streamlit.cache_data
def read_file(file_path: str, convert_to_markdown: bool = False) -> str:
    """ Return the file contents in string format. If file ends in .docx will convert it to json and return"""
    file_name, file_extension = os.path.splitext(file_path)

    if convert_to_markdown:
        with open(file_path, mode='rb') as f:
            # results = mammoth.convert_to_markdown(f)
            results = mammoth.convert_to_html(f)
            contents = convert_content_to_markdown(results.value)
        # contents = results.value
    elif file_extension == ".docx":
        # read in a document
        my_doc = docx.Document(file_path)

        # Find any tables and replace with json strings
        tmp_file = convert_tables_to_json_in_tmp__file(my_doc)

        # coerce to JSON using the standard options

        # contents = simplify(my_doc)

        # contents = textract.parsers.process(file_path)
        # print("Extracting contents from: %s" % tmp_file)
        contents = textract.process(tmp_file).decode('utf-8')
        os.remove(tmp_file)

    else:
        with open(file_path, mode='r') as f:  # TODO: Make sure you want to open with rb option
            contents = f.read()

    return str(contents)


def read_files(file_paths: Union[str, List[str]], convert_to_markdown: bool = False) -> str:
    if isinstance(file_paths, str):
        # If a single string is provided, treat it as a file path
        return read_file(file_paths, convert_to_markdown)
    elif isinstance(file_paths, list):
        # If a list is provided, loop through each file path and read the file
        concatenated_content = ""
        for path in file_paths:
            file_content = read_file(path, convert_to_markdown)
            concatenated_content += file_content + "\n\n"  # You can customize the separator if needed
        return concatenated_content
    else:
        return "Invalid input. Please provide a string or a list of strings (file paths)."


def dict_to_markdown_table(data, headers):
    # Create the header row
    markdown_table = "| " + " | ".join(headers) + " |\n"
    markdown_table += "| " + " | ".join(["-" * len(header) for header in headers]) + " |\n"

    # Iterate over the dictionary items and add rows to the table
    for row_data in data:
        markdown_table += "| " + " | ".join([str(row_data.get(header, '')) for header in headers]) + " |\n"

    return markdown_table


def extract_and_read_zip(file_path: str, accepted_file_types: list[str]) -> dict:
    unacceptable_file_prefixes = ['._']
    students_data = {}

    if file_path.endswith('.zip'):

        # Open the zip file
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            # Iterate over each file in the zip archive
            for file_info in zip_ref.infolist():
                # Extract the file name and directory name
                file_name = os.path.basename(file_info.filename)
                directory_name = os.path.dirname(file_info.filename)

                # Check if the directory name represents a student folder
                folder_name_delimiter = ' - '
                if directory_name and folder_name_delimiter in directory_name:
                    student_name = directory_name.split(folder_name_delimiter)[1]

                    # Check if the file has an accepted file type
                    if file_name.endswith(tuple(accepted_file_types)) and not file_name.startswith(
                            tuple(unacceptable_file_prefixes)):
                        # Read the file contents
                        with zip_ref.open(file_info.filename) as file:
                            # TODO: Change to modules on read file method
                            # file_contents = file.read().decode('utf-8')  # Assuming UTF-8 encoding
                            sub_file_name, sub_file_extension = os.path.splitext(file_name)
                            prefix = 'from_zip_' + str(randint(1000, 100000000)) + "_"
                            temp_file = tempfile.NamedTemporaryFile(delete=False, prefix=prefix,
                                                                    suffix=sub_file_extension)
                            temp_file.write(file.read())
                            # file_contents = read_file(
                            #    temp_file.name)  # Reading this way incase it may be .docx or some other type we want to pre-process differently

                        # Store the file contents in the dictionary
                        if student_name not in students_data:
                            students_data[student_name] = {}
                        # students_data[student_name][file_name] = file_contents
                        students_data[student_name][file_name] = temp_file.name

    return students_data


def login_if_needed(driver: WebDriver):
    # sleep for 3 seconds
    time.sleep(3)
    if "Web Login Service" in driver.title:
        # Duo Login
        duo_login(driver)
    elif "Sign in to your account" in driver.title:
        # Microsoft Login
        microsoft_login(driver)

def microsoft_login(driver: WebDriver):
    #Enter in user info and password
    instructor_user_id = os.environ["INSTRUCTOR_USERID"]
    instructor_password = os.environ["INSTRUCTOR_PASS"]

    wait = get_driver_wait(driver)

    original_window = driver.current_window_handle

    # Wait for title to change
    wait.until(EC.title_is("Sign in to your account"))

    # Wait for login elements
    wait.until(
        lambda d: d.find_element(By.XPATH, "//div[@id='loginHeader']"),
        "Waiting for login screen presence")

    # Enter username / email
    username_field = driver.find_element(By.NAME, "loginfmt")
    username_field.send_keys(instructor_user_id+"@cpcc.edu")
    # Click Next
    click_element_wait_retry(driver, wait, "//input[contains(@class, 'button_primary') and contains(@value,'Next')]", "Waiting for Next Button", By.XPATH)
    # Enter password
    password_field = driver.find_element(By.NAME, "passwd")
    password_field.send_keys(instructor_password)
    # Click Sign In
    click_element_wait_retry(driver, wait, "//input[contains(@class, 'button_primary') and contains(@value,'Sign in')]",
                             "Waiting for Sign in Button", By.XPATH)
    # Click the no button for Stay signed in
    no_stay_signed_in_button = click_element_wait_retry(driver, wait,
                                             "//input[contains(@class, 'button-secondary') and contains(@value,'No')]",
                                             "Waiting to click 'No' button to Stay Signed in")

    # Wait until login accepted
    wait.until(
        EC.invisibility_of_element(no_stay_signed_in_button),
        'Waiting for login to be successful')

    # Switch back to original window
    driver.switch_to.window(original_window)

def duo_login(driver: WebDriver):

    # TODO: This is not working when in streamlit cloud. Need to get values set before this line
    #from cqc_cpcc.utilities.env_constants import INSTRUCTOR_USERID, INSTRUCTOR_PASS

    instructor_user_id = os.environ["INSTRUCTOR_USERID"]
    instructor_password = os.environ["INSTRUCTOR_PASS"]

    wait = get_driver_wait(driver)

    original_window = driver.current_window_handle

    # Wait for title to change
    wait.until(EC.title_is("Web Login Service"))

    # Wait for login elements
    wait.until(
        lambda d: d.find_element(By.XPATH, "//div[@class='sr-only' and contains(text(),'Login')]"),
        "Waiting for login screen presence")

    # Login
    username_field = driver.find_element(By.ID, "username")
    password_field = driver.find_element(By.ID, "password")
    username_field.send_keys(instructor_user_id)
    password_field.send_keys(instructor_password)
    # login_field = driver.find_element(By.NAME, "_eventId_proceed")
    # login_field.click()
    click_element_wait_retry(driver, wait, "_eventId_proceed", "Waiting for login field", By.NAME)

    # Switch to Duo Iframe
    #duo_frame = wait.until(lambda d: d.find_element(By.ID, "duo_iframe"), "Waiting for Duo Iframe")
    #wait.until(EC.frame_to_be_available_and_switch_to_it(duo_frame))

    # NOTE: Duo push happens automatically now. Used to require a button push
    #click_element_wait_retry(driver, wait, "//button[contains(text(),'Send Me a Push')]", "Waiting for auth buttons")

    # Click the no to is this your device message
    login_message = click_element_wait_retry(driver, wait, "//button[contains(text(),'No, other people use this device')]", "Waiting to click 'No, other people use this device' button")

    # Wait until login accepted
    wait.until(
        EC.invisibility_of_element(login_message),
        'Waiting for login to be successful')

    # Switch back to original window
    driver.switch_to.window(original_window)

