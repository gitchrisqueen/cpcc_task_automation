from typing import List, Optional, Annotated

from docx import Document
from docx.shared import Pt
from langchain_core.callbacks import BaseCallbackHandler
from langchain_core.language_models import BaseChatModel
from pydantic import Field, BaseModel, StrictStr

from cqc_cpcc.utilities.AI.llm_deprecated.chains import get_exam_error_definition_from_completion_chain, \
    get_exam_error_definitions_completion_chain
from cqc_cpcc.utilities.AI.llm_deprecated.llms import get_default_llm
from cqc_cpcc.utilities.AI.exam_grading_openai import (
    DEFAULT_GRADING_MODEL,
    DEFAULT_TEMPERATURE,
    grade_exam_submission,
)
from cqc_cpcc.utilities.env_constants import SHOW_ERROR_LINE_NUMBERS
from cqc_cpcc.utilities.utils import ExtendedEnum, CodeError, ErrorHolder, merge_lists


def parse_error_type_enum_name(enum_name:str):
    parts = enum_name.split('_')  # Split the string by underscores
    course = parts[0]+" "+parts[1]  # First part is the course
    exam = parts[3]  # Second part is the exam
    name = '_'.join(parts[4:])  # Join the remaining parts to get the name
    return course, exam, name

class MajorErrorType(ExtendedEnum):
    """Enum representing various types of major coding errors."""

    """No documentation or insufficient documentation (comments)."""
    CSC_151_EXAM_1_INSUFFICIENT_DOCUMENTATION = "No documentation or insufficient amount of comments in the code"

    """Errors in coding sequence and selection."""
    # SEQUENCE_AND_SELECTION_ERROR = "Errors in the sequence and selection structures in the code"
    CSC_151_EXAM_1_SEQUENCE_AND_SELECTION_ERROR = "Errors in the coding sequence, selection and looping including incorrect use of comparison operators"

    CSC_151_EXAM_2_METHOD_ERRORS = """Method errors in the code (passing the incorrect number of arguments, incorrect data types for arguments and parameter variables, or failing to include the data type of parameter variables in the method header)"""

    """Errors that adversely impact the output (calculation errors, omissions, etc.)."""
    CSC_151_EXAM_1_OUTPUT_IMPACT_ERROR = "Errors that adversely impact the expected output, such as calculation errors or omissions"

    # CSC 251 Major Errors

    """No documentation or insufficient documentation (comments)."""
    CSC_251_EXAM_1_INSUFFICIENT_DOCUMENTATION = "No documentation or insufficient amount of comments in the code"

    """Errors in coding sequence and selection."""
    # SEQUENCE_AND_SELECTION_ERROR = "Errors in the sequence and selection structures in the code"
    CSC_251_EXAM_1_SEQUENCE_AND_SELECTION_ERROR = "Errors in the coding sequence, selection and looping including incorrect use of comparison operators"

    """Errors that adversely impact the output (calculation errors, omissions, etc.)."""
    CSC_251_EXAM_1_OUTPUT_IMPACT_ERROR = "Errors that adversely impact the expected output, such as calculation errors or omissions"

    """Failure to declare and use constants."""
    CSC_251_EXAM_1_CONSTANTS_ERROR = "Constants are not properly declared or used"

    """Decimal scale is not correct and/or missing comma separators."""
    CSC_251_EXAM_1_DECIMAL_SCALE = "There are issues with the expected code output where the decimal scale is not correct and/or missing commas separators"

    """Omit curly braces"""
    CSC_251_EXAM_1_CURLY_BRACES_OMITTED = "The code has omission of curly braces"

    """Security Holes"""
    CSC_251_EXAM_2_SECURITY_HOLES = "There are security holes in the code"

    """All method errors - adding or removing parameters, changing return type or name, returning incorrect values, etc."""
    CSC_251_EXAM_1_METHOD_ERRORS = "There are method errors (issues with parameters, return types, incorrect values, etc.)"

    """Any class design errors"""
    CSC_251_EXAM_1_CLASS_DESIGN_ERRORS = "There are class design errors"

    """Any ArrayList errors"""
    CSC_251_EXAM_1_ARRAYLIST_ERRORS = "There are errors involving ArrayList"

    """Any aggregation errors"""
    CSC_251_EXAM_2_AGGREGATION_ERRORS = "There are aggregation errors"

    # CSC 134 Project 1 Major Errors

    """Program does not compile"""
    CSC_134_PROJECT_1_DOES_NOT_COMPILE = "The program does not compile"

    """Major formatting issues – instructions are not followed for output (output and formatting must match)"""
    CSC_134_PROJECT_1_MAJOR_FORMATTING = "Major formatting issues – instructions are not followed for output (output and formatting must match)"

    """All errors when coding sequence, selection and/or looping control structures"""
    CSC_134_PROJECT_1_SEQUENCE_SELECTION_LOOPING_ERRORS = "There are errors in the coding sequence, selection and/or looping control structures"

    """Failure to validate input or input validation is incorrect"""
    CSC_134_PROJECT_1_INPUT_VALIDATION_ERRORS = "The code fails to validate input or input validation is incorrect"

    """Functions / Prototypes not declared and called correctly per the assignment instructions"""
    CSC_134_PROJECT_1_FUNCTION_PROTOTYPE_ERRORS = "Functions / Prototypes not declared and called correctly per the assignment instructions"

    """Calculation errors"""
    CSC_134_PROJECT_1_CALCULATION_ERRORS = "There are calculation errors"

    """Any errors that adversely impact the output"""
    CSC_134_PROJECT_1_OUTPUT_IMPACT_ERRORS = "There are errors that adversely impact the output"

    """No documentation or insufficient documentation (comments)"""
    CSC_134_PROJECT_2_INSUFFICIENT_DOCUMENTATION = "No documentation or insufficient amount of comments in the code"

    """Omit curly braces"""
    CSC_134_PROJECT_2_CURLY_BRACES_OMITTED = "The code has omission of curly braces"


    """Major Error Definitions for the CSC 152 Exam 1 (MideTerm)"""
    CSC_152_EXAM_1_STEP_ONE = "PROC SORT does not correctly sort by year and distance as specified"
    CSC_152_EXAM_1_STEP_TWO = "PROC TABULATE does not generate the required two-dimensional table"
    CSC_152_EXAM_1_STEP_THREE = "PROC REPORT does not correctly replicate the PROC TABULATE table"
    CSC_152_EXAM_1_STEP_FOUR = "PROC MEANS does not generate or display the required statistics"
    CSC_152_EXAM_1_STEP_FIVE = "PROC FREQ does not generate the required report with swimmer names and medal counts"
    CSC_152_EXAM_1_STEP_EIGHT = "The time is not displayed with two digits to the right of the decimal"

    CSC_152_EXAM_2_STEP_TWO = "The program does not satisfy the requirements in step 2"
    CSC_152_EXAM_2_STEP_EIGHT = "The program does not satisfy the requirements in step 8"

    """Major Error Definitions for the CSC 112 Assignment 1"""
    CSC_112_ASIGN_1_SIMPREGRES_RESULTS = "The regression results are missing"
    CSC_112_ASIGN_1_SIMPREGRES_FORMULA = "The regression formula is missing"
    CSC_112_ASIGN_1_SIMPREGRES_DIFFERENCE = "The explanation of difference is missing"


class MinorErrorType(ExtendedEnum):
    """Enum representing various types of minor coding errors."""

    """Syntax error in the code."""
    CSC_151_EXAM_1_SYNTAX_ERROR = "There are syntax errors in the code"

    """Violation of naming conventions."""
    CSC_151_EXAM_1_NAMING_CONVENTION = "Naming conventions are not followed"

    """Failure to declare and use constants."""
    CSC_151_EXAM_1_CONSTANTS_ERROR = "Constants are not properly declared or used"

    """Inefficient code practices."""
    CSC_151_EXAM_1_INEFFICIENT_CODE = "The code is inefficient and can be optimized"

    """Issues with the code output formatting."""
    CSC_151_EXAM_1_OUTPUT_FORMATTING = "There are issues with the expected code output formatting (spacing, decimal places, etc.)"

    """Programming style issues (indentation, white space, etc.)."""
    CSC_151_EXAM_1_PROGRAMMING_STYLE = "There are programming style issues that do not adhere to language standards (indentation, white space, etc.)"

    """Error related to the use of the Scanner class."""
    CSC_151_EXAM_1_SCANNER_CLASS = "There are errors related to the use of the Scanner class"

    # CSC 251 Minor Errors

    """Syntax error in the code."""
    CSC_251_EXAM_1_SYNTAX_ERROR = "There are syntax errors in the code"

    """Violation of naming conventions."""
    CSC_251_EXAM_1_NAMING_CONVENTION = "Naming conventions are not followed"

    """Inefficient code practices."""
    CSC_251_EXAM_1_INEFFICIENT_CODE = "The code is inefficient and can be optimized"

    """Programming style issues (indentation, white space, etc.)."""
    CSC_251_EXAM_1_PROGRAMMING_STYLE = "There are programming style issues that do not adhere to language standards (indentation, white space, etc.)"

    """Filename and class container name are not the same"""
    CSC_251_EXAM_1_FILE_CLASS_NAME_MISMATCH = "The filename and class container are not the same"

    """Minor formatting issues – not matching Sample Input and Output (ex: missing spaces, missing dollar sign, not using print/println appropriately)"""
    CSC_251_EXAM_1_MINOR_FORMATTING = "There are formatting issues not matching Sample Input and Output (i.e spacing, missing dollar sign, not using print/println appropriately, etc.)"

    """Stale data in classes"""
    CSC_251_EXAM_1_STALE_DATA = "There is stale data in classes"

    """Declaring variables/fields and not using them in the program"""
    CSC_251_EXAM_1_UNUSED_VARIABLES = "There are variables/fields declared that are not used in the program"

    """Incorrect data type used"""
    CSC_251_EXAM_1_INCORRECT_DATA_TYPE = "The program has the incorrect data type(s) used"

    """Program does not compile"""
    CSC_251_EXAM_1_DOES_NOT_COMPILE = "The program does not compile"


    # CSC 134 Project 1 Minor Errors

    """Incorrect filenames"""
    CSC_134_PROJECT_1_INCORRECT_FILENAMES = "There are incorrect filenames"

    """Filename and class container name are not the same"""
    CSC_134_PROJECT_2_FILE_CLASS_NAME_MISMATCH = "The filename and class container are not the same"

    """Minor formatting issues – not matching Sample Input and Output (ex: missing spaces, missing dollar sign, not using print/println appropriately)"""

    """Data type declaration is incorrect"""
    CSC_134_PROJECT_1_INCORRECT_DATA_TYPE_DECLARATION = "The code has data type declaration(s) that are incorrect"

    """Misspelling(s)"""
    CSC_134_PROJECT_1_MISSPELLINGS = "There are Misspelling(s)"

    """Failing to follow naming conventions"""
    CSC_134_PROJECT_1_NAMING_CONVENTION = "The code fails to follow naming conventions"

    """Failing to declare and use named constants (if needed)"""
    CSC_134_PROJECT_1_CONSTANTS_ERROR = "The code fails to declare and use named constants (if needed)"

    """Inefficient code (code duplication; spaghetti code, etc.)"""
    CSC_134_PROJECT_1_INEFFICIENT_CODE = "There is inefficient code (code duplication; spaghetti code, etc.)"

    """Minor formatting issues – decimal scale"""
    CSC_134_PROJECT_1_DECIMAL_SCALE = "Minor formatting issues – decimal scale"

    """Programming style (inconsistent or no indentation, inadequate white space, etc.)"""
    CSC_134_PROJECT_1_PROGRAMMING_STYLE = "Programming style (inconsistent or no indentation, inadequate white space, etc.)"

    #"""Minor Error Definitions for the CSC 152 Exam 1 (MideTerm)"""
    CSC_152_EXAM_1_STEP_SIX = "The dataset for 2020 medalists is not properly filtered"
    CSC_152_EXAM_1_STEP_SEVEN = "PROC FORMAT does not correctly replace rank with Gold, Silver, or Bronze"

    # """Minor Error Definitions for the CSC 152 Chapter 8 Exercise 23 Assignment"""
    CSC_152_ASIGN_CHP8EXC23_STEP_A = "Histogram does not accurately use the most recent population estimates for all countries"
    CSC_152_ASIGN_CHP8EXC23_STEP_B = "Histograms for each continent are not correctly separated or labeled"
    CSC_152_ASIGN_CHP8EXC23_STEP_C = "Box plot graph does not correctly group or display populations per continent"
    CSC_152_ASIGN_CHP8EXC23_STEP_D = "Comment does not adequately describe the statistical differences between histograms and box plots"
    CSC_152_ASIGN_CHP8EXC23_DATA_LABELS_ATTRIBUTES = "Program does not properly examine or report SAS data set labels and attributes"
    CSC_152_ASIGN_CHP8EXC23_HISTOGRAM_FORMAT = "Histograms are improperly formatted or missing required graphical elements"
    CSC_152_ASIGN_CHP8EXC23_BOXPLOT_FORMAT = "Box plots are improperly formatted or missing required graphical elements"
    CSC_152_ASIGN_CHP8EXC23_COMMENT_CLARITY = "Comment describing histogram and box plot differences lacks clarity or sufficient detail"
    CSC_152_EXAM_2_STEP_THREE = "The program does not satisfy the requirements in step 3"
    CSC_152_EXAM_2_STEP_FOUR = "The program does not satisfy the requirements in step 4"
    CSC_152_EXAM_2_STEP_FIVE = "The program does not satisfy the requirements in step 5"
    CSC_152_EXAM_2_STEP_SIX = "The program does not satisfy the requirements in step 6"
    CSC_152_EXAM_2_STEP_SEVEN = "The program does not satisfy the requirements in step 7"

    """MINOR Error Definitions for the CSC 112 Assignment 1"""
    CSC_112_ASIGN_1_MISSING_DATA_NOT_REMOVED = "The missing data was not removed"
    CSC_112_ASIGN_1_STATISTIC_NOT_IDENTIFIED = "The statistic was not identified"
    CSC_112_ASIGN_1_PREDICTED_VALUE_NOT_CALCULATED = "The predicted value was not calculated"

class MinorError(CodeError):
    """Class representing various types of minor coding errors."""
    error_type: Annotated[
        MinorErrorType,
        Field(description="The type of minor coding error.")
    ]


class MajorError(CodeError):
    """Class representing various types of major coding errors."""
    error_type: Annotated[
        MajorErrorType,
        Field(description="The type of major coding error.")
    ]


class JavaCode(BaseModel):
    """Class representing java code with functions to reference by line numbers"""

    entire_raw_code: StrictStr = Field(description="This raw string containing the entire Java code program")

    def get_line_number(self, line_of_code: StrictStr) -> List:
        line_number = None

        line_of_code = line_of_code.strip()

        # Escape the newlines characters
        # escaped_code = self.raw_code.encode('unicode_escape').decode('utf-8')
        # escaped_code = self.raw_code.replace("\n","\\\\n")
        escaped_code = self.entire_raw_code.strip()

        # print(escaped_code)

        # Split by remaining new lines
        code_lines = self.code_lines

        # print("Line count: %s" % str(len(code_lines)))

        # find line in text
        # if line_of_code in split_code:
        line_numbers = [code_lines.index(x) + 1 for x in code_lines if line_of_code in x]

        # print("Lines containing code: %s" % str(len(line_numbers)))

        if line_numbers is None:
            line_numbers = []

        # Return the line numbers
        return line_numbers

    @property
    def comments_count(self) -> int:
        return self.entire_raw_code.count("//") + self.entire_raw_code.count("/*")

    @property
    def code_lines(self) -> list[str]:
        return self.entire_raw_code.splitlines(keepends=True)

    @property
    def code_lines_count(self) -> int:
        return len(self.code_lines)

    @property
    def sufficient_amount_of_comments(self) -> bool:
        return (self.comments_count / self.code_lines_count) > .01


class ErrorDefinitions(ErrorHolder):
    """Class representing major and minor errors identified after doing a code analysis"""

    all_major_errors: Optional[List[MajorError]] = Field(description="List of the major code errors found in the code.")
    all_minor_errors: Optional[List[MinorError]] = Field(description="List of minor code errors found in the code.")

    def get_minor_errors_unique(self) -> None | List[MinorError]:
        minor_errors = None
        if self.all_minor_errors is not None:
            minor_errors = self.get_combined_errors_by_type(self.all_minor_errors)
        return minor_errors

    def get_major_errors_unique(self) -> None | List[MajorError]:
        major_errors = None
        if self.all_major_errors is not None:
            major_errors = self.get_combined_errors_by_type(self.all_major_errors)
        return major_errors


class CodeGrader:
    max_points: int
    major_errors: List[MajorError] = None
    minor_errors: List[MinorError] = None
    deduction_per_major_error: int
    deduction_per_minor_error: int
    error_definitions_completion_chain = None
    error_definitions_parser = None
    error_definitions_prompt = None
    
    # New OpenAI-based attributes
    exam_instructions: str = None
    exam_solution: str = None
    major_error_type_list: list = None
    minor_error_type_list: list = None
    use_openai_wrapper: bool = True  # Default to new implementation
    model_name: str = DEFAULT_GRADING_MODEL
    temperature: float = DEFAULT_TEMPERATURE
    use_openrouter: bool = False
    openrouter_auto_route: bool = True

    def __init__(self, max_points: int, exam_instructions: str, exam_solution: str,
                 deduction_per_major_error: int = 20,
                 deduction_per_minor_error: int = 5,
                 major_error_type_list: list = None,
                 minor_error_type_list: list = None,
                 grader_llm: BaseChatModel = None,
                 use_openai_wrapper: bool = True,
                 model_name: str = DEFAULT_GRADING_MODEL,
                 temperature: float = DEFAULT_TEMPERATURE,
                 use_openrouter: bool = False,
                 openrouter_auto_route: bool = True):
        self.max_points = max_points
        self.deduction_per_major_error = deduction_per_major_error
        self.deduction_per_minor_error = deduction_per_minor_error
        self.exam_instructions = exam_instructions
        self.exam_solution = exam_solution
        self.use_openai_wrapper = use_openai_wrapper
        self.model_name = model_name
        self.temperature = temperature
        self.use_openrouter = use_openrouter
        self.openrouter_auto_route = openrouter_auto_route
        
        if major_error_type_list is None:
            major_error_type_list = MajorErrorType.list()
        if minor_error_type_list is None:
            minor_error_type_list = MinorErrorType.list()
        
        self.major_error_type_list = major_error_type_list
        self.minor_error_type_list = minor_error_type_list

        # Keep LangChain setup for backward compatibility
        if not use_openai_wrapper:
            if grader_llm is None:
                grader_llm = get_default_llm()

            self.error_definitions_completion_chain, self.error_definitions_parser, self.error_definitions_prompt = get_exam_error_definitions_completion_chain(
                _llm=grader_llm,
                pydantic_object=ErrorDefinitions,
                major_error_type_list=major_error_type_list,
                minor_error_type_list=minor_error_type_list,
                exam_instructions=exam_instructions,
                exam_solution=exam_solution
            )

    @property
    def major_deduction_total_orig(self):
        return len(self.major_errors) * self.deduction_per_major_error

    @property
    def major_deduction_total(self):
        """ Calculates diminishing major error penalties as a geometric series: d * (1 - 0.5^n) / 0.5 """
        return self.deduction_per_major_error * (1 - 0.5 ** len(self.major_errors)) / (1 - 0.5)

    @property
    def minor_deduction_total(self):
        return len(self.minor_errors) * self.deduction_per_minor_error

    @property
    def total_deduction(self):
        return self.major_deduction_total + self.minor_deduction_total

    @property
    def points(self) -> float:
        # Ensure the calculated points are non-negative
        return max(self.max_points - self.total_deduction, 0)

    @property
    def final_score_text(self) -> str:
        return "Final Score: " + str(self.points) + " / " + str(self.max_points)

    @property
    def major_code_deduction_points_text(self) -> str:
        return "Major Code Errors: (-" + str(self.major_deduction_total) + " points)"

    @property
    def minor_code_deduction_points_text(self) -> str:
        return "Minor Code Errors: (-" + str(self.minor_deduction_total) + " points)"

    def get_text_feedback(self) -> str:
        grade_feedback = ""
        if self.major_errors is not None:
            grade_feedback += "\n" + self.major_code_deduction_points_text
            for error in self.major_errors:
                grade_feedback += "\n\n\t" + str(error).replace("\n", "\n\t")

        if self.minor_errors is not None:
            grade_feedback += "\n\n" + self.minor_code_deduction_points_text
            for error in self.minor_errors:
                grade_feedback += "\n\n\t" + str(error).replace("\n", "\n\t")

        grade_feedback += "\n\n" + self.final_score_text
        return grade_feedback

    async def grade_submission(self, student_submission: str, callback: BaseCallbackHandler = None):
        # print("Identifying Errors")
        
        if self.use_openai_wrapper:
            # New OpenAI wrapper path (supports OpenRouter)
            error_definitions = await grade_exam_submission(
                exam_instructions=self.exam_instructions,
                exam_solution=self.exam_solution,
                student_submission=student_submission,
                major_error_type_list=self.major_error_type_list,
                minor_error_type_list=self.minor_error_type_list,
                model_name=self.model_name,
                temperature=self.temperature,
                callback=callback,
                use_openrouter=self.use_openrouter,
                openrouter_auto_route=self.openrouter_auto_route,
            )
        else:
            # Legacy LangChain path
            error_definitions_from_llm = await get_exam_error_definition_from_completion_chain(
                student_submission=student_submission,
                completion_chain=self.error_definitions_completion_chain,
                parser=self.error_definitions_parser,
                prompt=self.error_definitions_prompt,
                callback=callback
            )
            error_definitions = ErrorDefinitions.model_validate(error_definitions_from_llm)
        
        # print("Errors Identified")

        # print("\n\nFinal Output:")
        # pprint(finalOutput)

        # print("\n\nCompiling Unique Errors")
        unique_major_errors = error_definitions.get_major_errors_unique()
        unique_minor_errors = error_definitions.get_minor_errors_unique()

        # Set major and minor errors to empty list if set to None
        if unique_major_errors is None:
            unique_major_errors = []
        if unique_minor_errors is None:
            unique_minor_errors = []

        # print("\n\nFinding Correct Error Line Numbers")
        jc = JavaCode(entire_raw_code=student_submission)

        # Remove or add Sufficient Commenting Error
        insufficient_comment_error = MajorError(error_type=MajorErrorType.CSC_151_EXAM_1_INSUFFICIENT_DOCUMENTATION,
                                                error_details="There is not enough comments to help others understand the purpose, functionality, and structure of the code.")

        if jc.sufficient_amount_of_comments:
            # print("Found Insufficient comments error by LLM but not true so removing")
            # Sufficient comments so remove this error type
            unique_major_errors = [x for x in unique_major_errors if
                                   x.error_type != MajorErrorType.CSC_151_EXAM_1_INSUFFICIENT_DOCUMENTATION]
        elif MajorErrorType.CSC_151_EXAM_1_INSUFFICIENT_DOCUMENTATION not in [x.error_type for x in unique_major_errors]:
            # print("Did not find Insufficient comments error by LLM but true so adding it")
            # Insufficient comments but error type doesnt exist
            unique_major_errors.insert(0, insufficient_comment_error)

        # Add the line numbers for where errors occur
        all_errors = merge_lists(unique_major_errors, unique_minor_errors)
        if SHOW_ERROR_LINE_NUMBERS and all_errors is not None:
            for code_error in all_errors:
                line_numbers = []
                if code_error.code_error_lines is not None:
                    for code_line in code_error.code_error_lines:
                        line_numbers_found = jc.get_line_number(code_line)
                        if line_numbers_found is not None:
                            line_numbers.extend(line_numbers_found)
                    # Return the unique list of line numbers
                    line_numbers = sorted(set(line_numbers))
                    code_error.set_line_numbers_of_error(line_numbers)

        # Append the errors back to the instance variables
        self.major_errors = unique_major_errors
        self.minor_errors = unique_minor_errors

    def add_errors_to_doc_document(self, d: Document, errors: List[CodeError]):
        for error in errors:
            # Add the error type
            type_paragraph = d.add_paragraph("", style='List Bullet 2')
            type_paragraph.add_run(str(error.error_type) + ":").italic = True

            # Add the error details
            for details in error.error_details.split("\n"):
                details_paragraph = d.add_paragraph("", style='List Bullet 3')
                details_paragraph.add_run(details)

    def set_document_style(self, document: Document):
        style = document.styles['Heading 3']
        font = style.font
        font.name = 'Lato'
        font.size = Pt(23)
        style2 = document.styles['List Bullet 2']
        font = style2.font
        font.name = 'Lato'
        font.size = Pt(19)
        style3 = document.styles['List Bullet 3']
        font = style3.font
        font.name = 'Lato'
        font.size = Pt(15)

    def save_feedback_to_docx(self, file_path: str):
        document = Document()

        # Set the styles for the document
        self.set_document_style(document)

        # Add the Major Errors, Deductions, and Details to the document
        if self.major_deduction_total > 0:
            document.add_heading(self.major_code_deduction_points_text, 3)
            self.add_errors_to_doc_document(document, self.major_errors)

        # Add the Minor Errors, Deductions, and Details to the document
        if self.minor_deduction_total > 0:
            document.add_heading(self.minor_code_deduction_points_text, 3)
            self.add_errors_to_doc_document(document, self.minor_errors)

        # Add The Final Score to the Document
        document.add_heading(self.final_score_text, 3)

        # Save the feedback to file
        document.save(file_path)
        # print("Feedback Saved to : %s" % file_path)

    def save_feedback_template(self, file_path: str):
        error_details = "This is an example detail for feedback."
        # Create all example errors
        self.major_errors = [MajorError(error_type=error_type, error_details=error_details) for error_type in
                             MajorErrorType.list()]
        self.minor_errors = [MinorError(error_type=error_type, error_details=error_details) for error_type in
                             MinorErrorType.list()]

        # print to file
        self.save_feedback_to_docx(file_path)


def start_grading():
    # Navigate to Bright Space

    # Look for Exams

    # For each Exam , get the instructions, the solution

    # Fore each student submission grade it

    # TODO: Put student name infront of output
    # points, grade_feedback = grade_submission(EXAM_INSTRUCTIONS, EXAM_SOLUTION, STUDENT_SUBMISSION)
    # print("Feedback:\n%s" % grade_feedback)

    # TODO: Add  grade output to submission as Draft

    don_nothing_for_now = None
