import datetime as DT
from typing import Optional, Annotated, List, TypeVar

from docx import Document
from docx.shared import Pt
from langchain_core.callbacks import BaseCallbackHandler
from langchain_core.language_models import BaseChatModel
from langchain_core.pydantic_v1 import Field
from pydantic.v1 import BaseModel
from selenium.common import TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait

from cqc_cpcc.exam_review import JavaCode
from cqc_cpcc.utilities.AI.llm.chains import get_feedback_completion_chain, \
    get_feedback_from_completion_chain
from cqc_cpcc.utilities.AI.llm.llms import get_default_llm
from cqc_cpcc.utilities.date import get_datetime
from cqc_cpcc.utilities.env_constants import *
from cqc_cpcc.utilities.logger import logger
from cqc_cpcc.utilities.selenium_util import get_session_driver, click_element_wait_retry, \
    get_elements_text_as_list_wait_stale, \
    get_elements_href_as_list_wait_stale
from cqc_cpcc.utilities.utils import are_you_satisfied, ExtendedEnum, CodeError, ErrorHolder, duo_login

COMMENTS_MISSING_STRING = "The code does not include sufficient commenting throughout"
T = TypeVar("T", bound=BaseModel)


def parse_error_type_enum_name(enum_name: str):
    parts = enum_name.split('_')  # Split the string by underscores
    course = parts[0] + " " + parts[1]  # First part is the course
    project = parts[3]  # Second part is the exam
    name = '_'.join(parts[4:])  # Join the remaining parts to get the name
    return course, project, name


class FeedbackType(ExtendedEnum):
    """Enum representing various types of feedback for code."""

    # General feedback for all programming languages
    COMMENTS_MISSING = "The code does not include sufficient commenting throughout"
    SYNTAX_ERROR = "There are syntax errors in the code"
    SPELLING_ERROR = "There are spelling mistakes in the code"
    OUTPUT_ALIGNMENT_ERROR = "There are output alignment issues in the code that will affect grades"
    PROGRAMMING_STYLE = "There are programming style issues that do not adhere to language standards"
    ADDITIONAL_TIPS_PROVIDED = "Additional insights regarding the code and learning"

    # Specific feedback for Java
    JAVA_NAMING_CONVENTION = "Naming conventions are not followed in the Java code"
    JAVA_CONSTANTS_ERROR = "Constants are not properly declared or used in the Java code"
    JAVA_INEFFICIENT_CODE = "The Java code is inefficient and can be optimized"
    JAVA_OUTPUT_FORMATTING = "There are issues with the expected Java code output formatting"
    JAVA_SCANNER_CLASS_ERROR = "There are errors related to the use of the Scanner class in Java"

    # Specific feedback for C++
    CPP_POINTER_ERROR = "There are errors related to the use of pointers in the C++ code"
    CPP_MEMORY_LEAK = "Potential memory leaks detected in the C++ code"
    CPP_SYNTAX_ERROR = "There are syntax errors specific to C++"
    CPP_NAMING_CONVENTION = "Naming conventions are not followed in the C++ code"
    CPP_CONSTANTS_ERROR = "Constants are not properly declared or used in the C++ code"

    # Future entry-level programming feedback
    VARIABLE_NAMING = "Variable names are not descriptive or do not follow naming conventions"
    FUNCTION_NAMING = "Function names are not descriptive or do not follow naming conventions"
    CODE_INDENTATION = "Code indentation is inconsistent or does not follow best practices"
    UNUSED_VARIABLES = "There are variables declared that are not used in the program"
    INCORRECT_DATA_TYPE = "The program has the incorrect data type(s) used"
    DOES_NOT_COMPILE = "The program does not compile"
    LOGIC_ERROR = "There are logical errors in the code that affect the program's functionality"
    MISSING_FUNCTIONALITY = "The code is missing required functionality as per the assignment instructions"


class DefaultFeedbackType(ExtendedEnum):
    """Enum representing various feedback types."""

    #### ------- Below is for JAVA expected submissions ------- #####
    """Checked for comments (indicated by // or /*) throughout the code."""
    CSC_151_PROJECT_ALL_COMMENTS_MISSING = COMMENTS_MISSING_STRING

    """Identified and addressed any syntax errors present."""
    CSC_151_PROJECT_ALL_SYNTAX_ERROR = "There are syntax errors in the code"

    """Scrutinized for spelling mistakes in the code."""
    CSC_151_PROJECT_ALL_SPELLING_ERROR = "There are spelling mistakes in the code"

    """Referenced any Output Alignment differences and emphasized its importance."""
    CSC_151_PROJECT_ALL_OUTPUT_ALIGNMENT_ERROR = "There are output alignment issues in the code that will affect exam grades"

    """Evaluated the programming style for adherence to language standards."""
    CSC_151_PROJECT_ALL_PROGRAMMING_STYLE = "There are programming style issues that do not adhere to java language standards"

    """Offered additional insights, knowledge, or tips."""
    CSC_151_PROJECT_ALL_ADDITIONAL_TIPS_PROVIDED = "Helpful insights regarding the submission to enhance learning"

    #### ------- Below is for DOCX expected submissions ------- #####
    # """Feedback to reference the expected input"""
    # EXPECTED_INPUTS = "There are missing variables or their corresponding data types in the Input column to ensure proper storage of user-entered values"
    # """Feedback to reference an ideal process"""
    # ALGORITHM_LOGIC = "The Process column does not clearly indicate the connection between the expected inputs from the Input column and desired output from the Output column"
    # """Feedback to reference expected output"""
    # EXPECTED_OUTPUT = "The Output column fails to show the desired outputs or maintain proper formatting for the desired outputs"
    # """Feedback to share additional insight regarding their submission as related to the instructions and example solution"""
    # FINAL_THOUGHTS = "Additional insight regarding possible improvements"


class Feedback(CodeError):
    """Class representing various types of feedback for coding."""
    # DefaultFeedbackType,
    error_type: Annotated[
        FeedbackType,
        Field(description="The type of feedback for code.")
    ]


class FeedbackGuide(ErrorHolder):
    """Class representing feedback after reviewing a student's submission"""
    all_feedback: Optional[List[Feedback]] = Field(description="List of all the feedback for the submission.")

    def get_feedback_unique(self) -> None | List[Feedback]:
        feedback_errors = None
        if self.all_feedback is not None:
            feedback_errors = self.get_combined_errors_by_type(self.all_feedback)
        return feedback_errors



def init_page(driver: WebDriver, wait: WebDriverWait) -> str:
    driver.get(BRIGHTSPACE_URL)
    # Wait for title to change
    wait.until(EC.title_is("Web Login Service"))

    original_window = driver.current_window_handle

    # Login
    duo_login(driver)

    # Wait for title to change
    wait.until(EC.title_is(BRIGHTSPACE_HOMEPAGE_TITLE))

    # Return the original_window
    return original_window


def give_project_feedback():
    driver, wait = get_session_driver()

    # original_window = init_page(driver, wait)

    # Find All Courses in Brightspace
    course_urls = get_course_urls(driver, wait)
    logger.info("Course Urls:\n%s" % "\n".join(course_urls))

    # Process submissions for each course
    for url in course_urls:
        process_submissions_from_course_url(driver, wait, url)

    logger.info("Feedback Complete")


def get_course_urls(driver: WebDriver, wait: WebDriverWait) -> list[str]:
    # Click on the course menu
    click_element_wait_retry(driver, wait, "d2l-navigation-s-course-menu", "Waiting for Course Menu", By.CLASS_NAME)

    # Get the course url
    course_links = wait.until(
        lambda d: d.find_elements(By.XPATH,
                                  "//div[contains(@class,'d2l-course-selector-item')]/a[contains(@class,'d2l-link') and not(contains(text(),'Sandbox'))]"),
        "Waiting for Course Links")
    course_urls = map(lambda b: b.get_attribute("href"), course_links)

    return list(course_urls)


def process_submissions_from_course_url(driver: WebDriver, wait: WebDriverWait, url: str):
    handles = driver.window_handles

    # Opens a new tab and switches to new tab
    driver.switch_to.new_window('tab')

    # Wait for the new window or tab
    wait.until(EC.new_window_is_opened(handles))

    # Keep track of current tab
    # current_tab = driver.current_window_handle

    # Navigate to course url
    driver.get(url)

    # Navigate to Course Tools -> Assignments
    assignments_link = wait.until(
        lambda d: d.find_element(By.XPATH,
                                 "//d2l-menu-item-link[@text='Assignments' and contains(@class,'d2l-navigation-s-menu-item')]"),
        'Waiting for Assignments link')
    assignments_url = BRIGHTSPACE_URL + assignments_link.get_attribute('href')
    logger.info("Assignments URL: %s", assignments_url)

    driver.get(assignments_url)

    # Find Projects with Past Due Dates
    # Criteria - Name contains "Project", Has Submissions
    try:
        due_dates = get_elements_text_as_list_wait_stale(wait,
                                                         "//tr[descendant::a[contains(@title,'Project')] and descendant::a[contains(@class,'d2l-link') and contains(@title,'new')]]//*[contains(text(),'Due')]",
                                                         "Waiting on Due Dates")
    except TimeoutException as te:
        logger.info("No projects with new submissions. Returning...")
        return

    due_dates = list(map(lambda b: b.replace("Due on ", ""), due_dates))
    logger.info("Due Dates:\n%s" % "\n".join(due_dates))

    submission_links = get_elements_href_as_list_wait_stale(wait,
                                                            "//tr[descendant::a[contains(@title,'Project')] and descendant::a[contains(@class,'d2l-link') and contains(@title,'new')]]//a[contains(@class,'d2l-link') and contains(@title,'new')]",
                                                            "Waiting on submission links")

    logger.info("Submission Links:\n%s" % "\n".join(submission_links))

    due_date_submission_links_dict = dict(zip(due_dates, submission_links))
    logger.info("Due Date Submission Links:\n%s" % str(due_date_submission_links_dict))

    shrt_frmt = "%m/%d"
    today = DT.date.today()
    today_string = today.strftime(shrt_frmt)
    yesterday = today - DT.timedelta(days=1)
    # yesterday_string = yesterday.strftime(shrt_frmt)
    week_ago = yesterday - DT.timedelta(days=7)
    week_ago_string = week_ago.strftime(shrt_frmt)

    for due_date in due_date_submission_links_dict:
        proper_date = get_datetime(due_date)
        proper_date_str = proper_date.strftime(shrt_frmt)
        if week_ago <= proper_date.date() <= today:
            logger.info("%s IS between %s - %s" % (proper_date_str, week_ago_string, today_string))
            submission_link = due_date_submission_links_dict[due_date]
            logger.info("Submission Link: %s" % submission_link)

            process_submission_from_link(driver, wait, submission_link)
        else:
            logger.info("%s is NOT between %s - %s" % (proper_date_str, week_ago_string, today_string))

    # TODO: Give feedback to missing submissions

    are_you_satisfied()
    # Close the tab
    driver.close()


def process_submission_from_link(driver: WebDriver, wait: WebDriverWait, url: str):
    handles = driver.window_handles

    # Opens a new tab and switches to new tab
    driver.switch_to.new_window('tab')

    # Wait for the new window or tab
    wait.until(EC.new_window_is_opened(handles))

    # Keep track of current tab
    # current_tab = driver.current_window_handle

    # Navigate to course url
    driver.get(url)

    # Loop until header numbers match i.e 17 - 17

    # Download and parse the .Java Files

    # TODO: Leave feedback

    # TODO: Add Feedback signature to end of feedback
    # FEEDBACK_SIGNATURE

    # Save as Draft !!!

    are_you_satisfied()
    # Close the tab
    driver.close()




class FeedbackGiver:
    feedback_list: List[Feedback] = None
    feedback_completion_chain = None
    feedback_parser = None
    feedback_prompt = None
    feedback_guide: FeedbackGuide = None

    def __init__(self,
                 course_name: str,
                 assignment_instructions: str,
                 assignment_solution: str,
                 wrap_code_in_markdown: bool = True,
                 feedback_llm: BaseChatModel = None,
                 feedback_type_list: list = None,

                 ):
        if feedback_llm is None:
            feedback_llm = get_default_llm()

        self.feedback_completion_chain, self.feedback_parser, self.feedback_prompt = get_feedback_completion_chain(
            llm=feedback_llm,
            pydantic_object=FeedbackGuide,
            feedback_type_list=feedback_type_list,
            assignment=assignment_instructions,
            solution=assignment_solution,
            course_name=course_name,
            wrap_code_in_markdown=wrap_code_in_markdown
        )

    async def generate_feedback(self, student_submission: str, callback: BaseCallbackHandler = None):
        # print("Identifying Errors")
        feedback_from_llm = await get_feedback_from_completion_chain(
            student_submission=student_submission,
            completion_chain=self.feedback_completion_chain,
            parser=self.feedback_parser,
            prompt=self.feedback_prompt,
            callback=callback

        )
        # print("\n\nFinal Output From LLM:")
        # pprint(feedback_from_llm)

        feedback_guide = FeedbackGuide.parse_obj(feedback_from_llm)

        unique_feedback = feedback_guide.get_feedback_unique()
        # Set feedback to empty list if set to None
        if unique_feedback is None:
            unique_feedback = []


        if COMMENTS_MISSING_STRING in DefaultFeedbackType.list():

            # print("\n\nFinding Correct Error Line Numbers")
            jc = JavaCode(entire_raw_code=student_submission)

            if feedback_guide.all_feedback is not None:
                unique_feedback = feedback_guide.get_feedback_unique()

                for code_error in unique_feedback:
                    line_numbers = []
                    if code_error.code_error_lines is not None:
                        for code_line in code_error.code_error_lines:
                            line_numbers_found = jc.get_line_number(code_line)
                            if line_numbers_found is not None:
                                line_numbers.extend(line_numbers_found)
                        # Return the unique list of line numbers
                        line_numbers = sorted(set(line_numbers))
                        code_error.set_line_numbers_of_error(line_numbers)

                # Create Insufficient Comments if applicable
                insufficient_comment_error = Feedback(
                    error_type=DefaultFeedbackType.CSC_151_PROJECT_ALL_COMMENTS_MISSING,
                    error_details="There is not enough comments to help others understand the purpose, functionality, and structure of the code.")

                if jc.sufficient_amount_of_comments:
                    # print("Found Insufficient comments error by LLM but not true so removing")
                    # Sufficient comments so remove this error type
                    unique_feedback = [x for x in unique_feedback if
                                       x.error_type != DefaultFeedbackType.CSC_151_PROJECT_ALL_COMMENTS_MISSING]
                elif DefaultFeedbackType.CSC_151_PROJECT_ALL_COMMENTS_MISSING not in [x.error_type for x in
                                                                                      unique_feedback]:
                    # print("Did not find Insufficient comments error by LLM but true so adding it")
                    # Insufficient comments but error type doesnt exist
                    unique_feedback.insert(0, insufficient_comment_error)

        # Append the feedback back to the instance variables
        self.feedback_list = unique_feedback

    def set_document_style(self, document: Document):
        style = document.styles['Heading 3']
        font = style.font
        font.name = 'Lato'
        font.size = Pt(23)
        style2 = document.styles['List Bullet 2']
        font = style2.font
        font.name = 'Lato'
        font.size = Pt(19)
        style3 = document.styles['List Bullet 3']
        font = style3.font
        font.name = 'Lato'
        font.size = Pt(15)

    def add_feedback_to_doc_document(self, d: Document, errors: List[CodeError]):
        for error in errors:
            # Add the error type
            type_paragraph = d.add_paragraph("", style='List Bullet 2')
            type_paragraph.add_run(str(error.error_type) + ":").italic = True

            # Add the error details
            for details in error.error_details.split("\n"):
                details_paragraph = d.add_paragraph("", style='List Bullet 3')
                details_paragraph.add_run(details)

    def save_feedback_to_docx(self, file_path: str, pre_text: str = "", post_text: str = "",
                              pre_post_heading_size: int = 3):
        document = Document()

        # Set the styles for the document
        self.set_document_style(document)

        feedback_list = self.feedback_list

        # Add the pre-text to the document
        if len(pre_text) > 0:
            document.add_heading(pre_text, pre_post_heading_size)

        # Add the feedback to the document
        if len(feedback_list) > 0:
            self.add_feedback_to_doc_document(document, feedback_list)

        # Add the pre-text to the document
        if len(post_text) > 0:
            document.add_heading(post_text, pre_post_heading_size)

        # Save the feedback to file
        document.save(file_path)
        print("Feedback Saved to : %s" % file_path)
