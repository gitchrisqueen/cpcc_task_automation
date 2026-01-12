#  Copyright (c) 2024. Christopher Queen Consulting LLC (http://www.ChristopherQueenConsulting.com/)

"""CPCC-branded Word document generator for student feedback.

This module generates polished Word documents containing student feedback
following CPCC branding guidelines. Documents are formatted for student
consumption and do NOT include numeric scores, percentages, or grade bands.

CPCC Branding Spec:
- Colors: CPCC Blue (#005AA3), CPCC Light Blue (#7BAFD4)
- Font: Calibri
- Structure: Title, Student info, Course/Assignment, Divider, Sections
- Sections: Summary, Strengths, Improvements, Errors Observed (if applicable)
"""

import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from cqc_cpcc.rubric_models import RubricAssessmentResult
from cqc_cpcc.student_feedback_builder import build_student_feedback


# CPCC Brand Colors
CPCC_BLUE = RGBColor(0, 90, 163)  # #005AA3
CPCC_LIGHT_BLUE = RGBColor(123, 175, 212)  # #7BAFD4
NEUTRAL_DARK = RGBColor(31, 41, 55)  # #1F2937
NEUTRAL_GRAY = RGBColor(107, 114, 128)  # #6B7280
ACCENT_GRAY = RGBColor(209, 213, 219)  # #D1D5DB


def normalize_color_to_hex(color) -> str:
    """Normalize various color input formats to a 6-character uppercase hex string.
    
    This function accepts multiple color input formats and converts them to a
    consistent hex string format suitable for use in Word document XML.
    
    Supported input formats:
    - python-docx RGBColor objects (which are tuples of (r, g, b))
    - tuple/list of (r, g, b) values (0-255 range)
    - hex string with or without '#' prefix ("AABBCC" or "#AABBCC")
    
    Args:
        color: Color in one of the supported formats
        
    Returns:
        6-character uppercase hex string without '#' prefix (e.g., "005AA3")
        
    Raises:
        ValueError: If color format is not recognized or values are out of range
        TypeError: If color type is not supported
        
    Examples:
        >>> normalize_color_to_hex(RGBColor(0x12, 0x34, 0x56))
        '123456'
        >>> normalize_color_to_hex((18, 52, 86))
        '123456'
        >>> normalize_color_to_hex([18, 52, 86])
        '123456'
        >>> normalize_color_to_hex("#123456")
        '123456'
        >>> normalize_color_to_hex("123456")
        '123456'
    """
    # Handle RGBColor objects (which are tuples) and regular tuples/lists
    if isinstance(color, (RGBColor, tuple, list)):
        if len(color) != 3:
            raise ValueError(f"Color tuple/list must have exactly 3 values (r, g, b), got {len(color)}")
        
        r, g, b = color
        
        # Validate range
        for component, name in [(r, 'r'), (g, 'g'), (b, 'b')]:
            if not isinstance(component, int) or not (0 <= component <= 255):
                raise ValueError(
                    f"Color component '{name}' must be an integer between 0 and 255, got {component}"
                )
        
        return f"{r:02X}{g:02X}{b:02X}"
    
    # Handle hex string
    elif isinstance(color, str):
        # Remove '#' prefix if present
        hex_str = color.lstrip('#')
        
        # Validate hex string length
        if len(hex_str) != 6:
            raise ValueError(
                f"Hex color string must be 6 characters (without '#'), got '{hex_str}' ({len(hex_str)} chars)"
            )
        
        # Validate hex characters
        try:
            int(hex_str, 16)
        except ValueError:
            raise ValueError(f"Invalid hex color string: '{hex_str}' contains non-hex characters")
        
        # Return uppercase
        return hex_str.upper()
    
    else:
        raise TypeError(
            f"Unsupported color type: {type(color).__name__}. "
            f"Expected RGBColor, tuple/list (r,g,b), or hex string."
        )


def sanitize_filename(name: str) -> str:
    """Sanitize a string for use in a filename.
    
    Removes or replaces characters that are invalid in filenames:
    - Removes slashes, colons, asterisks, question marks, quotes, pipes
    - Replaces spaces with underscores
    - Handles unicode edge cases
    
    Args:
        name: Original name string
        
    Returns:
        Sanitized filename-safe string
    """
    # Remove invalid filename characters
    sanitized = re.sub(r'[<>:"/\\|?*]', '', name)
    # Replace spaces with underscores
    sanitized = sanitized.replace(' ', '_')
    # Remove leading/trailing periods, spaces, and underscores
    sanitized = sanitized.strip('. _')
    # Ensure it's not empty (after stripping all the above)
    if not sanitized:
        sanitized = "Feedback"
    return sanitized


def add_horizontal_line(paragraph, color: RGBColor = CPCC_LIGHT_BLUE, height: int = 15000):
    """Add a horizontal line (divider) to a paragraph.
    
    Args:
        paragraph: docx Paragraph object
        color: RGB color for the line (RGBColor, tuple, or hex string)
        height: Line thickness in EMUs (15000 = ~1.5pt)
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(height // 1000))  # Convert to eighths of a point
    bottom.set(qn('w:space'), '1')
    
    # Normalize color to hex string using our utility function
    color_hex = normalize_color_to_hex(color)
    bottom.set(qn('w:color'), color_hex)
    
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_student_feedback_doc(
    student_name: str,
    course_id: str,
    assignment_name: str,
    feedback_result: RubricAssessmentResult,
    metadata: Optional[dict] = None
) -> bytes:
    """Generate a CPCC-branded Word document containing student feedback.
    
    Creates a polished Word document following CPCC branding guidelines.
    The document contains:
    - Title: "Feedback Summary"
    - Student info line
    - Course/Assignment/Date line
    - Horizontal divider
    - Summary section
    - Strengths (bulleted)
    - Improvements (bulleted)
    - Errors Observed (sub-bullets per error, only if errors exist)
    
    Document does NOT include:
    - Numeric points or percentages
    - Grade bands or rubric scores
    - Internal error codes
    
    Args:
        student_name: Student's name
        course_id: Course identifier (e.g., "CSC151")
        assignment_name: Assignment name (e.g., "Exam 1")
        feedback_result: RubricAssessmentResult from grading
        metadata: Optional metadata dict (timestamp, instructor name, etc.)
        
    Returns:
        bytes: DOCX file content as bytes
    """
    # Create document
    doc = Document()
    
    # Set default font and margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Title: "Feedback Summary" (20pt Bold, CPCC Blue)
    title = doc.add_heading('Feedback Summary', level=1)
    title_run = title.runs[0]
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = CPCC_BLUE
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0
    
    # Student line (12pt Regular, Neutral Dark)
    student_para = doc.add_paragraph()
    student_run = student_para.add_run(f'Student: {student_name}')
    student_run.font.name = 'Calibri'
    student_run.font.size = Pt(12)
    student_run.font.color.rgb = NEUTRAL_DARK
    student_para.paragraph_format.space_after = Pt(6)
    
    # Course/Assignment/Date line (12pt Regular, Neutral Dark)
    date_str = metadata.get('date') if metadata else datetime.now().strftime('%Y-%m-%d')
    course_para = doc.add_paragraph()
    course_run = course_para.add_run(
        f'Course: {course_id} | Assignment: {assignment_name} | Date: {date_str}'
    )
    course_run.font.name = 'Calibri'
    course_run.font.size = Pt(12)
    course_run.font.color.rgb = NEUTRAL_DARK
    course_para.paragraph_format.space_after = Pt(12)
    
    # Horizontal divider (CPCC Light Blue)
    divider_para = doc.add_paragraph()
    add_horizontal_line(divider_para, CPCC_LIGHT_BLUE)
    divider_para.paragraph_format.space_after = Pt(12)
    
    # Build student feedback text (no scores)
    feedback_text = build_student_feedback(
        feedback_result,
        student_name=student_name,
        include_greeting=False  # Skip greeting in doc
    )
    
    # Parse feedback text into sections
    sections_dict = _parse_feedback_sections(feedback_text)
    
    # Add Summary section
    if 'summary' in sections_dict and sections_dict['summary']:
        _add_section_heading(doc, 'Summary')
        summary_para = doc.add_paragraph(sections_dict['summary'])
        _format_body_paragraph(summary_para)
    
    # Add Strengths section
    if 'strengths' in sections_dict and sections_dict['strengths']:
        _add_section_heading(doc, 'Strengths')
        for strength in sections_dict['strengths']:
            para = doc.add_paragraph(strength, style='List Bullet')
            _format_body_paragraph(para)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
    
    # Add Improvements section
    if 'improvements' in sections_dict and sections_dict['improvements']:
        _add_section_heading(doc, 'Areas for Improvement')
        for improvement in sections_dict['improvements']:
            para = doc.add_paragraph(improvement, style='List Bullet')
            _format_body_paragraph(para)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
    
    # Add Errors Observed section (only if errors exist)
    if 'errors' in sections_dict and sections_dict['errors']:
        _add_section_heading(doc, 'Errors Observed')
        
        # Group errors by severity
        major_errors = [e for e in sections_dict['errors'] if e.get('severity') == 'major']
        minor_errors = [e for e in sections_dict['errors'] if e.get('severity') == 'minor']
        
        if major_errors:
            major_label = doc.add_paragraph('Major Issues:')
            major_label_run = major_label.runs[0]
            major_label_run.font.italic = True
            _format_body_paragraph(major_label)
            major_label.paragraph_format.space_after = Pt(4)
            
            for error in major_errors:
                _add_error_to_doc(doc, error)
        
        if minor_errors:
            minor_label = doc.add_paragraph('Minor Issues:')
            minor_label_run = minor_label.runs[0]
            minor_label_run.font.italic = True
            _format_body_paragraph(minor_label)
            minor_label.paragraph_format.space_after = Pt(4)
            
            for error in minor_errors:
                _add_error_to_doc(doc, error)
    
    # Convert document to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()


def _add_section_heading(doc: Document, heading_text: str):
    """Add a section heading (14pt Bold, CPCC Blue).
    
    Args:
        doc: Document object
        heading_text: Heading text
    """
    heading = doc.add_heading(heading_text, level=2)
    heading_run = heading.runs[0]
    heading_run.font.name = 'Calibri'
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.color.rgb = CPCC_BLUE
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.line_spacing = 1.0


def _format_body_paragraph(paragraph):
    """Format a paragraph with body text styling (11pt, Neutral Dark).
    
    Args:
        paragraph: Paragraph object
    """
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = NEUTRAL_DARK
    
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)


def _parse_feedback_sections(feedback_text: str) -> dict:
    """Parse structured feedback text into sections.
    
    Extracts:
    - summary: Opening paragraph(s)
    - strengths: List items under "Strengths"
    - improvements: List items under "Areas for Improvement"
    - errors: Structured error info under "Errors Observed"
    
    Args:
        feedback_text: Formatted feedback text from build_student_feedback
        
    Returns:
        Dict with keys: summary, strengths, improvements, errors
    """
    sections = {
        'summary': '',
        'strengths': [],
        'improvements': [],
        'errors': []
    }
    
    lines = feedback_text.split('\n')
    current_section = 'summary'
    current_error = None
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip empty lines
        if not line_stripped:
            continue
        
        # Detect section headers
        if line_stripped.startswith('**Strengths:**'):
            current_section = 'strengths'
            continue
        elif line_stripped.startswith('**Areas for Improvement:**'):
            current_section = 'improvements'
            continue
        elif line_stripped.startswith('**Errors Observed:**'):
            current_section = 'errors'
            continue
        elif line_stripped.startswith('*Major Issues:*') or line_stripped.startswith('*Minor Issues:*'):
            # Track severity for errors
            if 'Major' in line_stripped:
                current_section = 'errors_major'
            else:
                current_section = 'errors_minor'
            continue
        
        # Process content based on current section
        if current_section == 'summary':
            sections['summary'] += line_stripped + ' '
        elif current_section == 'strengths':
            if line_stripped.startswith('- '):
                sections['strengths'].append(line_stripped[2:])
        elif current_section == 'improvements':
            if line_stripped.startswith('- '):
                sections['improvements'].append(line_stripped[2:])
        elif current_section in ('errors_major', 'errors_minor'):
            if line_stripped.startswith('- **'):
                # New error entry
                # Format: "- **Error Name**: Description"
                match = re.match(r'-\s*\*\*(.+?)\*\*:\s*(.+)', line_stripped)
                if match:
                    error_name = match.group(1)
                    error_desc = match.group(2)
                    current_error = {
                        'name': error_name,
                        'description': error_desc,
                        'severity': 'major' if current_section == 'errors_major' else 'minor',
                        'notes': ''
                    }
                    sections['errors'].append(current_error)
            elif current_error and line_stripped and not line_stripped.startswith('Keep up'):
                # Additional notes for current error
                current_error['notes'] += line_stripped + ' '
    
    # Clean up summary
    sections['summary'] = sections['summary'].strip()
    
    return sections


def _add_error_to_doc(doc: Document, error: dict):
    """Add an error entry to the document with sub-bullets.
    
    Format:
    - Error Title (human readable)
      - Description
      - Notes (if available)
    
    Args:
        doc: Document object
        error: Error dict with name, description, severity, notes
    """
    # Main error bullet
    error_para = doc.add_paragraph(error['name'], style='List Bullet')
    _format_body_paragraph(error_para)
    error_para.runs[0].font.bold = True
    error_para.paragraph_format.space_after = Pt(2)
    
    # Description sub-bullet
    desc_para = doc.add_paragraph(error['description'], style='List Bullet 2')
    _format_body_paragraph(desc_para)
    desc_para.paragraph_format.left_indent = Inches(0.5)
    desc_para.paragraph_format.space_after = Pt(2)
    
    # Notes sub-bullet (if available)
    if error.get('notes') and error['notes'].strip():
        notes_para = doc.add_paragraph(error['notes'].strip(), style='List Bullet 2')
        _format_body_paragraph(notes_para)
        notes_para.paragraph_format.left_indent = Inches(0.5)
        notes_para.paragraph_format.space_after = Pt(6)
    else:
        # Add spacing after description if no notes
        desc_para.paragraph_format.space_after = Pt(6)
